import os
import re
from pathlib import Path
from .base import BuddyPlugin


class DOCXPlugin(BuddyPlugin):
    name = "docx"
    supported_extensions = ["docx", "doc"]
    description = "Read and edit Word documents"

    def execute(self, file_path: str, task: str) -> str:
        path = self.validate_file(file_path)
        try:
            from docx import Document
        except ImportError:
            return "❌ python-docx not installed. Run: pip install python-docx"

        lo = task.lower()
        if any(k in lo for k in ["read", "extract", "text", "show", "content"]):
            return self._read(path, Document)
        if any(k in lo for k in ["info", "metadata", "meta"]):
            return self._info(path, Document)
        if "heading" in lo:
            return self._add_heading(path, task, Document)
        if any(k in lo for k in ["append", "add text", "add paragraph", "add line"]):
            return self._append(path, task, Document)
        if any(k in lo for k in ["replace", "find"]):
            return self._replace(path, task, Document)
        return self._read(path, Document)

    def _read(self, path: Path, Document) -> str:
        doc = Document(str(path))
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        if not text:
            return "❌ Document appears empty."
        if len(text) > 3000:
            return f"📝 Content (first 3000 chars):\n\n{text[:3000]}…"
        return f"📝 Content:\n\n{text}"

    def _info(self, path: Path, Document) -> str:
        doc = Document(str(path))
        words = sum(len(p.text.split()) for p in doc.paragraphs)
        return (
            f"📝 {path.name}\n"
            f"   Paragraphs: {len(doc.paragraphs)}\n"
            f"   Words:      ~{words}\n"
            f"   Sections:   {len(doc.sections)}\n"
            f"   Size:       {os.path.getsize(str(path)) // 1024} KB"
        )

    def _add_heading(self, path: Path, task: str, Document) -> str:
        m = re.search(r'heading[:\s]+(.+)', task, re.I)
        if not m:
            return "❌ Format: heading: My Title"
        doc = Document(str(path))
        doc.add_heading(m.group(1).strip(), level=1)
        doc.save(str(path))
        return f"✅ Heading '{m.group(1).strip()}' added to {path.name}"

    def _append(self, path: Path, task: str, Document) -> str:
        m = re.search(r'(?:append|add)[:\s]+(.+)', task, re.I | re.DOTALL)
        if not m:
            return "❌ Format: append: your new text here"
        doc = Document(str(path))
        doc.add_paragraph(m.group(1).strip())
        doc.save(str(path))
        return f"✅ Text appended to {path.name}"

    def _replace(self, path: Path, task: str, Document) -> str:
        m = re.search(r'replace\s+"(.+?)"\s+with\s+"(.+?)"', task, re.I)
        if not m:
            m = re.search(r'replace\s+(.+?)\s+with\s+(.+)', task, re.I)
        if not m:
            return '❌ Format: replace "old text" with "new text"'
        old, new = m.group(1).strip(), m.group(2).strip()
        doc = Document(str(path))
        count = 0
        for para in doc.paragraphs:
            for run in para.runs:
                if old in run.text:
                    run.text = run.text.replace(old, new)
                    count += 1
        doc.save(str(path))
        return f"✅ Replaced {count} instance(s) of '{old}' → '{new}' in {path.name}"
